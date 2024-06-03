# import streamlit as st
# import pdfplumber
# from docx import Document
# import pythoncom
# import win32com.client
# from PIL import Image
# from transformers import pipeline
# import os

# # Initialize Hugging Face pipelines
# qa_pipeline = pipeline("question-answering")
# summarizer = pipeline("summarization")

# # Function to convert PDF to text
# def pdf_to_text(file):
#     with pdfplumber.open(file) as pdf:
#         text = ""
#         for page in pdf.pages:
#             extracted_text = page.extract_text()
#             if extracted_text:
#                 text += extracted_text + "\n"
#     return text

# # Function to convert text to DOCX
# def text_to_docx(text, output_path):
#     doc = Document()
#     for line in text.split("\n"):
#         doc.add_paragraph(line)
#     doc.save(output_path)

# # Function to convert DOC to DOCX using pywin32
# def doc_to_docx(input_path, output_path):
#     pythoncom.CoInitialize()
#     word = win32com.client.Dispatch('Word.Application')
#     word.Visible = False
#     doc = word.Documents.Open(input_path)
#     doc.SaveAs(output_path, FileFormat=16)  # 16 corresponds to wdFormatXMLDocument (DOCX)
#     doc.Close()
#     word.Quit()

# # Function to convert DOC to PDF
# def doc_to_pdf(input_path, output_path):
#     try:
#         pythoncom.CoInitialize()
#         word = win32com.client.Dispatch('Word.Application')
#         word.Visible = False
#         doc = word.Documents.Open(input_path)
#         doc.SaveAs(output_path, FileFormat=17)  # 17 corresponds to wdFormatPDF
#         doc.Close()
#         word.Quit()
#     except Exception as e:
#         st.error(f"An error occurred while converting DOC to PDF: {e}")

# # Function to perform question answering
# def ask_question(question, context):
#     result = qa_pipeline(question=question, context=context)
#     return result['answer']

# # Function to summarize text
# def summarize_text(text):
#     try:
#         summaries = summarizer(text, max_length=150, min_length=30, do_sample=False)
#         if summaries:
#             return summaries[0]['summary_text']
#         else:
#             return "Summarization failed or returned empty."
#     except Exception as e:
#         return str(e)

# # Streamlit app
# st.title("File Conversion and Q&A App")

# # Sidebar menu
# menu = ["Text Files", "Image Files", "Ask Questions from PDF", "Summarize PDF"]
# choice = st.sidebar.selectbox("Select File Type", menu)

# if choice == "Text Files":
#     st.header("Text File Conversions")
#     text_menu = ["PDF to DOCX", "DOC to PDF"]
#     text_choice = st.sidebar.selectbox("Select an option", text_menu)

#     if text_choice == "PDF to DOCX":
#         st.subheader("Convert PDF to DOCX")
#         uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
        
#         if uploaded_file is not None:
#             # Convert PDF to text
#             text = pdf_to_text(uploaded_file)
            
#             if text:
#                 # Convert text to DOCX
#                 output_path = "converted.docx"
#                 text_to_docx(text, output_path)
                
#                 # Provide download link for DOCX file
#                 with open(output_path, "rb") as file:
#                     st.download_button(
#                         label="Download DOCX",
#                         data=file,
#                         file_name="converted.docx",
#                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                     )
#             else:
#                 st.error("No text found in PDF")

#     elif text_choice == "DOC to PDF":
#         st.subheader("Convert DOC to PDF")
#         uploaded_file = st.file_uploader("Choose a DOC file", type="doc")
        
#         if uploaded_file is not None:
#             # Save the uploaded DOC file
#             doc_path = "uploaded.doc"
#             with open(doc_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())

#             # Convert DOC to PDF
#             output_path = os.path.join(os.getcwd(), "converted.pdf")
#             doc_to_pdf(doc_path, output_path)
            
#             # Provide download link for PDF file
#             with open(output_path, "rb") as file:
#                 st.download_button(
#                     label="Download PDF",
#                     data=file,
#                     file_name="converted.pdf",
#                     mime="application/pdf"
#                 )

# elif choice == "Image Files":
#     st.header("Image File Conversions")
#     image_menu = ["PNG to JPG", "JPG to PNG", "PNG to PDF", "JPG to PDF"]
#     image_choice = st.sidebar.selectbox("Select an option", image_menu)

#     if image_choice == "PNG to JPG":
#         st.subheader("Convert PNG to JPG")
#         uploaded_file = st.file_uploader("Choose a PNG file", type="png")
        
#         if uploaded_file is not None:
#             # Save the uploaded PNG file
#             png_path = "uploaded.png"
#             with open(png_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())

#             # Convert PNG to JPG
#             output_path = "converted.jpg"
#             image = Image.open(png_path)
#             # Ensure image is in RGB mode before saving as JPEG
#             image = image.convert("RGB")
#             image.save(output_path, "JPEG")
            
#             # Provide download link for JPG file
#             with open(output_path, "rb") as file:
#                 st.download_button(
#                     label="Download JPG",
#                     data=file,
#                     file_name="converted.jpg",
#                     mime="image/jpeg"
#                 )
    
#     elif image_choice == "JPG to PNG":
#         st.subheader("Convert JPG to PNG")
#         uploaded_file = st.file_uploader("Choose a JPG file", type="jpg")
        
#         if uploaded_file is not None:
#             # Save the uploaded JPG file
#             jpg_path = "uploaded.jpg"
#             with open(jpg_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())

# elif choice == "Ask Questions from PDF":
#     st.header("Ask Questions from PDF")
#     uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
#     if uploaded_file is not None:
#         # Convert PDF to text
#         context = pdf_to_text(uploaded_file)
        
#         st.text_area("Extracted Text from PDF", context, height=300)
        
#         question = st.text_input("Ask a question about the PDF")
        
#         if st.button("Get Answer"):
#             if question:
#                 answer = ask_question(question, context)
#                 st.text_area("Answer", answer, height=100)
#             else:
#                 st.error("Please enter a question.")

# elif choice == "Summarize PDF":
#     st.header("Summarize PDF")
#     uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
#     if uploaded_file is not None:
#         # Convert PDF to text
#         text = pdf_to_text(uploaded_file)
        
#         if text:
#             # Ensure text length is appropriate for summarization
#             if len(text) > 1000:
#                 text = text[:1000]  # Truncate text if too long
            
#             summary = summarize_text(text)
#             st.text_area("Summary", summary, height=200)
#         else:
#             st.error("No text found in PDF")
import streamlit as st
import pdfplumber
from docx import Document
import pythoncom
import win32com.client
from PIL import Image
from transformers import pipeline
import os

# Initialize Hugging Face pipelines
qa_pipeline = pipeline("question-answering")
summarizer = pipeline("summarization")

# Function to convert PDF to text
def pdf_to_text(file):
    with pdfplumber.open(file) as pdf:
        text = ""
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
    return text

# Function to convert text to DOCX
def text_to_docx(text, output_path):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)

# Function to convert DOC to DOCX using pywin32
def doc_to_docx(input_path, output_path):
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=16)  # 16 corresponds to wdFormatXMLDocument (DOCX)
    doc.Close()
    word.Quit()

# Function to convert DOC to PDF
def doc_to_pdf(input_path, output_path):
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=17)  # 17 corresponds to wdFormatPDF
        doc.Close()
        word.Quit()
    except Exception as e:
        st.error(f"Error converting DOC to PDF: {e}")

# Function to perform question answering
def ask_question(question, context):
    result = qa_pipeline(question=question, context=context)
    return result['answer']

# Function to summarize text
def summarize_text(text):
    try:
        summaries = summarizer(text, max_length=150, min_length=30, do_sample=False)
        if summaries:
            return summaries[0]['summary_text']
        else:
            return "Summarization failed or returned empty."
    except Exception as e:
        return str(e)

# Streamlit app
st.title("File Conversion and Q&A App")

# Sidebar menu
menu = ["Text Files", "Image Files", "Ask Questions from PDF", "Summarize PDF"]
choice = st.sidebar.selectbox("Select File Type", menu)

if choice == "Text Files":
    st.header("Text File Conversions")
    text_menu = ["PDF to DOCX", "DOC to PDF"]
    text_choice = st.sidebar.selectbox("Select an option", text_menu)

    if text_choice == "PDF to DOCX":
        st.subheader("Convert PDF to DOCX")
        uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
        
        if uploaded_file is not None:
            # Convert PDF to text
            text = pdf_to_text(uploaded_file)
            
            if text:
                # Convert text to DOCX
                output_path = "converted.docx"
                text_to_docx(text, output_path)
                
                # Provide download link for DOCX file
                with open(output_path, "rb") as file:
                    st.download_button(
                        label="Download DOCX",
                        data=file,
                        file_name="converted.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error("No text found in PDF")

    elif text_choice == "DOC to PDF":
        st.subheader("Convert DOC to PDF")
        uploaded_file = st.file_uploader("Choose a DOC file", type="doc")
        
        if uploaded_file is not None:
            # Save the uploaded DOC file
            doc_path = "uploaded.doc"
            with open(doc_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            # Convert DOC to PDF
            output_path = "converted.pdf"
            doc_to_pdf(doc_path, output_path)
            
            # Provide download link for PDF file
            with open(output_path, "rb") as file:
                st.download_button(
                    label="Download PDF",
                    data=file,
                    file_name="converted.pdf",
                    mime="application/pdf"
                )

elif choice == "Image Files":
    st.header("Image File Conversions")
    image_menu = ["PNG to JPG", "JPG to PNG", "PNG to PDF", "JPG to PDF"]
    image_choice = st.sidebar.selectbox("Select an option", image_menu)

    if image_choice == "PNG to JPG":
        st.subheader("Convert PNG to JPG")
        uploaded_file = st.file_uploader("Choose a PNG file", type="png")
        
        if uploaded_file is not None:
            # Save the uploaded PNG file
            png_path = "uploaded.png"
            with open(png_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            # Convert PNG to JPG
            output_path = "converted.jpg"
            image = Image.open(png_path)
            # Ensure image is in RGB mode before saving as JPEG
            image = image.convert("RGB")
            image.save(output_path, "JPEG")
            
            # Provide download link for JPG file
            with open(output_path, "rb") as file:
                st.download_button(
                    label="Download JPG",
                    data=file,
                    file_name="converted.jpg",
                    mime="image/jpeg"
                )
    
    elif image_choice == "JPG to PNG":
        st.subheader("Convert JPG to PNG")
        uploaded_file = st.file_uploader("Choose a JPG file", type="jpg")
        
        if uploaded_file is not None:
            # Save the uploaded JPG file
            jpg_path = "uploaded.jpg"
            with open(jpg_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

elif choice == "Ask Questions from PDF":
    st.header("Ask Questions from PDF")
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
    if uploaded_file is not None:
        # Convert PDF to text
        context = pdf_to_text(uploaded_file)
        
        st.text_area("Extracted Text from PDF", context, height=300)
        
        question = st.text_input("Ask a question about the PDF")
        
        if st.button("Get Answer"):
            if question:
                answer = ask_question(question, context)
                st.text_area("Answer", answer, height=100)
            else:
                st.error("Please enter a question.")

elif choice == "Summarize PDF":
    st.header("Summarize PDF")
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
    
    if uploaded_file is not None:
        # Convert PDF to text
        text = pdf_to_text(uploaded_file)
        
        if text:
            # Ensure text length is appropriate for summarization
            if len(text) > 1000:
                text = text[:1000]  # Truncate text if too long
            
            summary = summarize_text(text)
            st.text_area("Summary", summary, height=200)
        else:
            st.error("No text found in PDF")


